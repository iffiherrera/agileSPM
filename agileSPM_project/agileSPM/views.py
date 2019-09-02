from django.shortcuts import render, reverse, redirect
from django.http import HttpResponse, HttpResponseRedirect, HttpResponseNotFound
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import SOWScrum, SOWKanban, SOWScrumban, User
from .forms import SOWScrumForm, RegisterForm, SOWKanbanForm, SOWScrumbanForm
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
import datetime
from django.urls import reverse_lazy
from django.views.generic.edit import DeleteView

def contact(request):
    return render(request, 'agileSPM/base_agileSPM.html',{})

def register(request):

    if request.method == 'POST':
        register_form = RegisterForm(request.POST)
        if register_form.is_valid():
            register_form.save()
            username = register_form.cleaned_data.get('username')
            email = register_form.cleaned_data.get('email')
            password1 = register_form.cleaned_data.get('password1')
            password2 = register_form.cleaned_data.get('password2')
            print('username was succesful', username)
            return redirect('my_docs')
    else:
        register_form = RegisterForm()

    return render(request,'agileSPM/registration/register.html', {'register_form' : register_form})

# Home view 
def index(request):
    user_form = RegisterForm()
    context_dict = {
        'user': request.user,
        'user_form': user_form,
    }
    return render(request, 'agileSPM/index.html', context=context_dict)

# Scrum form 
@login_required
def scrumForm(request):
    form = SOWScrumForm()

    if request.method == 'POST':
        form = SOWScrumForm(data=request.POST or None)

        if form.is_valid(): 
            form = form.save(commit=False)
            form.author = request.user
            print('user', request.user)
            form.save()
            id = form.id
            print('id created:', id)
            return redirect('success_scrum', id=id)

        else:
            print(form.errors)
    
    else:
        form = SOWScrumForm()

    context_dict = {'user': request.user,
                    'form': form}

    return render(request,'agileSPM/scrum/full_form.html', context=context_dict)

# Successful completion of form view.
def success_scrum(request, id):
    complete_scrum_forms = SOWScrum.objects.filter(author=request.user)
    context_dict = {'id': id,
                    'complete_scrum_forms': complete_scrum_forms,}

    return render(request,'agileSPM/done.html', context=context_dict)


def scrumForm_update(request, id):
    instance = SOWScrum.objects.get(id=id)
    form = SOWScrumForm(data=request.POST or None, instance=instance)
    edited_scrum_forms = SOWScrum.objects.filter(author=request.user)

    if form.is_valid():
        instance = form.save(commit=False)
        form.save()
        # id = form.id
        print('id edited:',id)
        return redirect('success_scrum', id=id)

    context_dict = {
        'id': id,
        'instance': instance,
        'scrum_form': edited_scrum_forms,
        'form':form}

    return render(request, 'agileSPM/scrum/full_form.html', context=context_dict)
   
class DeleteScrumForm(DeleteView):
    model = SOWScrum
    template_name = 'agileSPM/delete_doc.html'
    success_url = reverse_lazy('my_docs')
           
@login_required
def kanbanForm(request):
    form = SOWKanbanForm()

    if request.method == 'POST':
        form = SOWKanbanForm(data=request.POST)

        if form.is_valid(): 
            form = form.save(commit=False)
            form.author = request.user
            print('user', request.user)
            form.save()
            id = form.id
            print('id created:', id)
            return redirect('success_kanban', id=id)

        else:
            print(form.errors)
    
    else:
        form = SOWKanbanForm()
    
    context_dict = {'user': request.user,
                    'form': form}

    return render(request,'agileSPM/kanban/kanban_form.html', context=context_dict)

# Successful completion of form view.
def success_kanban(request, id):
    complete_kanban_forms = SOWKanban.objects.filter(author=request.user)
    context_dict = {'id': id,
                    'complete_scrum_forms': complete_kanban_forms}

    return render(request,'agileSPM/kanban/done_kanban.html', context=context_dict)

def kanbanForm_update(request, id):
    instance = SOWKanban.objects.get(id=id)
    form = SOWKanbanForm(data=request.POST or None, instance=instance)
    edited_kanban_forms = SOWScrum.objects.filter(author=request.user)
    if form.is_valid():
        instance = form.save(commit=False)
        form.save()
        # id = form.id
        print('id edited:',id)
        return redirect('success_kanban', id=id)

    context_dict = {
        'id': id,
        'instance': instance,
        'kanban_form': edited_kanban_forms,
        'form':form}

    return render(request, 'agileSPM/kanban/kanban_form.html', context=context_dict)

class DeleteKanbanForm(DeleteView):
    model = SOWKanban
    template_name = 'agileSPM/delete_doc.html'
    success_url = reverse_lazy('my_docs')

@login_required
def scrumbanForm(request):
    form = SOWScrumbanForm()

    if request.method == 'POST':
        form = SOWScrumbanForm(data=request.POST)

        if form.is_valid(): 
            form = form.save(commit=False)
            form.author = request.user
            print('user', request.user)
            form.save()
            id = form.id
            print('id created:', id)
            return redirect('success_scrumban', id=id)

        else:
            print(form.errors)
    
    else:
        form = SOWScrumbanForm()
    
    context_dict = {'user': request.user,
                    'form': form}

    return render(request,'agileSPM/scrumban/scrumban_form.html', context=context_dict)

# Successful completion of form view.
def success_scrumban(request, id):
    complete_scrumban_forms = SOWScrumban.objects.filter(author=request.user)
    context_dict = {'id': id,
                    'complete_scrumban_forms': complete_scrumban_forms}

    return render(request,'agileSPM/scrumban/done_scrumban.html', context=context_dict)

def scrumbanForm_update(request, id):
    instance = SOWScrumban.objects.get(id=id)
    form = SOWScrumbanForm(data=request.POST or None, instance=instance)
    edited_scrumban_forms = SOWScrumban.objects.filter(author=request.user)
    if form.is_valid():
        instance = form.save(commit=False)
        form.save()
        # id = form.id
        print('id edited:',id)
        return redirect('success_scrumban', id=id)

    context_dict = {
        'id': id,
        'instance': instance,
        'scrumban_form': edited_scrumban_forms,
        'form':form}

    return render(request, 'agileSPM/scrumban/scrumban_form.html', context=context_dict)

class DeleteScrumbanForm(DeleteView):
    model = SOWScrumban
    template_name = 'agileSPM/delete_doc.html'
    success_url = reverse_lazy('my_docs')

# User account view 
@login_required
def my_docs(request):

    user_scrum_forms = SOWScrum.objects.filter(author=request.user)
    user_kanban_forms = SOWKanban.objects.filter(author=request.user)
    user_scrumban_forms = SOWScrumban.objects.filter(author=request.user)

    print('scrum',user_scrum_forms)
    print('kanban',user_kanban_forms)
    print('scrumban',user_scrumban_forms)

    context_dict = {'user': request.user,
                    'id': id,
                    'user_scrum_forms': user_scrum_forms,
                    'user_kanban_forms': user_kanban_forms,
                    'user_scrumban_forms': user_scrumban_forms }
   
    print('user',request.user)

    return render(request, 'agileSPM/docs.html', context=context_dict)
  
# Scrum Document creation using docx-Python API
def scrum_doc(request, id):
    user_input = SOWScrum.objects.get(id=id)
    
## Document formatting, rules & styles ##
    s_project = Document()
    paragraph = s_project.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.keep_with_next = True
 
    styles = s_project.styles
    body_text_style = s_project.styles['Body Text']
    list_bullet_style = s_project.styles['List Bullet']
    numbered_bullet_style = s_project.styles['ListNumber']

    # Casting datetime fields as string
    date = str(user_input.date_project)
    milestones = str(user_input.milestones)
    delivery = str(user_input.delivery)
    date_signature1 = str(user_input.date_signature1)
    date_signature2 = str(user_input.date_signature2)

    # Casting int values to string
    sprints = str(user_input.sprint)
    length = str(user_input.sprintLength)
    invoice = str(user_input.invoice)

## Cover page ##

    # Title
    s_project.add_heading(user_input.title, 0)
    # print(s_project)
    # Produced by 
    s_project.add_paragraph(user_input.produced_by, style=body_text_style)
    # Date created
    s_project.add_paragraph(date, style=body_text_style)
    s_project.add_page_break()

## Overview ##

    # Introduction
    s_project.add_heading('Introduction', level=1)
    s_project.add_paragraph(user_input.intro, style=body_text_style)

## Objectives ##

    # Deliverables 
    s_project.add_heading('Deliverables',level=1)
    s_project.add_paragraph(user_input.deliverables, style=list_bullet_style)
    # Assumptions
    s_project.add_heading('Assumptions',level=1)
    s_project.add_paragraph(user_input.assumptions, style=list_bullet_style)

## Scope ##

    # In Scope
    s_project.add_heading('In Scope',level=1)
    s_project.add_paragraph(user_input.outScope, style=body_text_style)
    # Out of Scope
    s_project.add_heading('Out of Scope',level=1)
    s_project.add_paragraph(user_input.inScope, style=body_text_style)

## Backlog ##

    # Product backlog 
    s_project.add_heading('Product backlog',level=1)
    s_project.add_paragraph(user_input.backlog, style=list_bullet_style)

## Mode of Delivery ##

    # Sprints
    para = s_project.add_paragraph('Number of Sprints:')
    para.add_run(sprints).bold = True
    # Sprint Length
    para = s_project.add_paragraph('Length of Sprint:')
    para.add_run(length).bold = True
    # Team 
    s_project.add_heading('Team Structure',level=1)
    s_project.add_paragraph(user_input.team, style=body_text_style)
    # Done
    s_project.add_heading('Defining Done',level=1)
    s_project.add_paragraph(user_input.done, style=body_text_style)
    # Review
    s_project.add_heading('Review Opportunities',level=1)
    s_project.add_paragraph(user_input.review, style=body_text_style)
  

## Milestones ## 

    # Milestones table 
    s_project.add_heading('Milestones',level=1)
    s_project.add_paragraph(milestones, style=list_bullet_style)
    s_project.add_paragraph(user_input.milestone_description, style=body_text_style)


    #### Code awaiting input from milestone model formset, commented out as was not able to complete it in time. #####

    # milestone_table = s_project.add_table(rows=3, cols=2)
    # milestone_table.style = 'LightShading-Accent1'
    # header_cells = milestone_table.rows[0].cells
    # header_cells[0].text = 'Milestone'
    # header_cells[1].text = 'Date'
    # row = milestone_table.add_row() # Method adding a row, to be called! 
    # Delivery
    para = s_project.add_paragraph('Delivery date:')
    para.add_run(delivery).bold = True

    # for row in milestone_table.rows:
    #     for cell in row.cells:
    #         print(cell.text)

## Costs ## 

    # Invoice schedule table
    s_project.add_heading('Invoice schedule',level=1)
    s_project.add_paragraph(invoice, style=body_text_style)
    s_project.add_paragraph(user_input.invoice_info, style=body_text_style)


    #### Code awaiting input from invoice model formset, commented out as was not able to complete it in time. #####

    # invoice_table = s_project.add_table(rows=3, cols=3)
    # invoice_table.style = 'LightShading-Accent1'
    # header_cells = invoice_table.rows[0].cells
    # header_cells[0].text = 'Invoice number'
    # header_cells[1].text = 'Date'
    # header_cells[2].text = 'Total'
    # row = invoice_table.add_row() # Method adding a row, to be called! 

    # for row in invoice_table.rows:
    #         for cell in row.cells:
    #             print(cell.text)

## Acceptance ##

    # Signature 1
    s_project.add_heading('Acceptance',level=1)
    para = s_project.add_paragraph('Full Name:')
    para.add_run(user_input.firstName).bold = True
    s_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = s_project.add_paragraph('Date signed:')
    para.add_run(date_signature1)
    # Signature 2
    para = s_project.add_paragraph('Full Name:')
    para.add_run(user_input.secondName).bold = True
    s_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = s_project.add_paragraph('Date signed:')
    para.add_run(date_signature2)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=ScrumProject.docx'
    s_project.save(response)
    return response

    return render(request)

# Kanban Document creation using docx-Python API
def kanban_doc(request, id):
    user_input = SOWKanban.objects.get(id=id)

## Document formatting,rules & styles ##
    k_project = Document()
    paragraph = k_project.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.keep_with_next = True

    styles = k_project.styles
    body_text_style = k_project.styles['Body Text']
    list_bullet_style = k_project.styles['List Bullet']
    numbered_bullet_style = k_project.styles['ListNumber']

    # Casting datetime fields as string
    date = str(user_input.date_project)
    delivery = str(user_input.delivery)
    date_signature1 = str(user_input.date_signature1)
    date_signature2 = str(user_input.date_signature2)

    # Casting int values to string
    columns = str(user_input.columns)
    wipLimit = str(user_input.wipLimit)
    invoice = str(user_input.invoice)

## Cover page ##

    # Title
    k_project.add_heading(user_input.title, 0)
    # Produced by 
    k_project.add_paragraph(user_input.produced_by, style=body_text_style)
    # Date created
    k_project.add_paragraph(date, style=body_text_style)
    k_project.add_page_break()

## Overview ##

    # Introduction
    k_project.add_heading('Introduction', level=1)
    k_project.add_paragraph(user_input.intro, style=body_text_style)

## Objectives ##

    # Deliverables 
    k_project.add_heading('Deliverables',level=1)
    k_project.add_paragraph(user_input.deliverables, style=list_bullet_style)
    # Assumptions
    k_project.add_heading('Assumptions',level=1)
    k_project.add_paragraph(user_input.assumptions, style=list_bullet_style)

## Scope ##

    # In Scope
    k_project.add_heading('In Scope',level=1)
    k_project.add_paragraph(user_input.inScope, style=body_text_style)
    # Out of Scope
    k_project.add_heading('Out of Scope',level=1)
    k_project.add_paragraph(user_input.outScope, style=body_text_style)

## Backlog ##

    # Product backlog 
    k_project.add_heading('Product backlog',level=1)
    k_project.add_paragraph(user_input.backlog, style=list_bullet_style)

## Mode of Delivery ##

    # Kanban plan 
    k_project.add_heading('Kanban Overview',level=1)
    k_project.add_paragraph(user_input.plan, style=body_text_style)
    # Columns
    para = k_project.add_paragraph('Number of Columns used in board:')
    para.add_run(columns).bold = True
    # Labels
    k_project.add_paragraph(user_input.column_labels, style=body_text_style)
    # WIP limit
    para = k_project.add_paragraph('Work In Progress limit:')
    para.add_run(wipLimit).bold = True
    # Delivery
    para = k_project.add_paragraph('Delivery date:')
    para.add_run(delivery).bold = True

## Costs ## 

    # Invoice Amount
    k_project.add_heading('Invoice amount',level=1)
    k_project.add_paragraph(invoice, style=body_text_style)
    k_project.add_paragraph(user_input.invoice_info, style=body_text_style)
    
    #### Code awaiting input from invoice model formset, commented out as was not able to complete it in time. #####

    # invoice_table = k_project.add_table(rows=3, cols=3)
    # invoice_table.style = 'LightShading-Accent1'
    # header_cells = invoice_table.rows[0].cells
    # header_cells[0].text = 'Invoice number'
    # header_cells[1].text = 'Date'
    # header_cells[2].text = 'Total'
    # row = invoice_table.add_row() # Method adding a row, to be called! 

    # for row in table.rows:
    #         for cell in rows.cells:
    #             print(cell.text)

## Acceptance ##

    # Signature 1
    k_project.add_heading('Acceptance',level=1)
    para = k_project.add_paragraph('Full Name:')
    para.add_run(user_input.firstName).bold = True
    k_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = k_project.add_paragraph('Date signed:')
    para.add_run(date_signature1)
    # Signature 2
    para = k_project.add_paragraph('Full Name:')
    para.add_run(user_input.secondName).bold = True
    k_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = k_project.add_paragraph('Date signed:')
    para.add_run(date_signature2)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=KanbanProject.docx'
    k_project.save(response)
    return response

    return render(request)

# Scrumban Document creation using docx-Python API
def scrumban_doc(request,id):
    user_input = SOWScrumban.objects.get(id=id)

## Document formatting, styles and rules ##
    sb_project = Document()

    paragraph = sb_project.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.keep_with_next = True

    styles = sb_project.styles
    body_text_style = sb_project.styles['Body Text']
    list_bullet_style = sb_project.styles['List Bullet']
    numbered_bullet_style = sb_project.styles['ListNumber']

    # Casting datetime fields as string
    date = str(user_input.date_project)
    delivery = str(user_input.delivery)
    milestones = str(user_input.milestones)
    date_signature1 = str(user_input.date_signature1)
    date_signature2 = str(user_input.date_signature2)

    # Casting int values to string
    wipLimit = str(user_input.wipLimit)
    invoice = str(user_input.invoice)
    iterations = str(user_input.iterations)

## Cover page ##

    # Title
    sb_project.add_heading(user_input.title, 0)
    # Produced by 
    sb_project.add_paragraph(user_input.produced_by, style=body_text_style)
    # Date created
    sb_project.add_paragraph(date, style=body_text_style)
    sb_project.add_page_break()

## Overview ##

    # Introduction
    sb_project.add_heading('Introduction', level=1)
    sb_project.add_paragraph(user_input.intro, style=body_text_style)

## Objectives ##

    # Deliverables 
    sb_project.add_heading('Deliverables',level=1)
    sb_project.add_paragraph(user_input.deliverables, style=list_bullet_style)
    # Assumptions
    sb_project.add_heading('Assumptions',level=1)
    sb_project.add_paragraph(user_input.assumptions, style=list_bullet_style)

## Scope ##

    # In Scope
    sb_project.add_heading('In Scope',level=1)
    sb_project.add_paragraph(user_input.inScope, style=body_text_style)
    # Out of Scope
    sb_project.add_heading('Out of Scope',level=1)
    sb_project.add_paragraph(user_input.outScope, style=body_text_style)

## Backlog ##

    # Product backlog 
    sb_project.add_heading('Product backlog',level=1)
    sb_project.add_paragraph(user_input.backlog, style=list_bullet_style)

## Mode of Delivery ##

    # Scrumban plan 
    sb_project.add_heading('Scrumban Overview',level=1)
    sb_project.add_paragraph(user_input.plan, style=body_text_style)
    # Iterations
    para = sb_project.add_paragraph('Number of Iterations:')
    para.add_run(iterations).bold = True
    # WIP Limit
    para = sb_project.add_paragraph('Work In Progress Limit:')
    para.add_run(wipLimit).bold = True
    # Team 
    sb_project.add_heading('Team Structure',level=1)
    sb_project.add_paragraph(user_input.team, style=body_text_style)
    # Review
    sb_project.add_heading('Review Opportunities',level=1)
    sb_project.add_paragraph(user_input.review, style=body_text_style)
  

## Milestones ## 

    # Milestones table 
    sb_project.add_heading('Milestones',level=1)
    sb_project.add_paragraph(milestones, style=list_bullet_style)
    sb_project.add_paragraph(user_input.milestone_description, style=body_text_style)

    #### Code awaiting input from milestone model formset, commented out as was not able to complete it in time. #####

    # milestone_table = sb_project.add_table(rows=3, cols=2)
    # milestone_table.style = 'LightShading-Accent1'
    # header_cells = milestone_table.rows[0].cells
    # header_cells[0].text = 'Milestone'
    # header_cells[1].text = 'Date'
    # row = milestone_table.add_row() # Method adding a row, to be called! 

    # for row in table.rows:
    #     for cell in rows.cells:
    #         print(cell.text)



   # Delivery
    para = sb_project.add_paragraph('Delivery date:')
    para.add_run(delivery).bold = True

## Costs ## 

    # Invoice amount
    sb_project.add_heading('Invoice schedule',level=1)
    sb_project.add_paragraph(invoice, style=body_text_style)
    sb_project.add_paragraph(user_input.invoice_info, style=body_text_style)

    #### Code awaiting input from invoice model formset, commented out as was not able to complete it in time. #####

    # invoice_table = sb_project.add_table(rows=3, cols=3)
    # invoice_table.style = 'LightShading-Accent1'
    # header_cells = invoice_table.rows[0].cells
    # header_cells[0].text = 'Invoice number'
    # header_cells[1].text = 'Date'
    # header_cells[2].text = 'Total'
    # row = invoice_table.add_row() # Method adding a row, to be called! 

    # for row in table.rows:
    #     for cell in rows.cells:
    #         print(cell.text)



## Acceptance ##

   # Signature 1
    sb_project.add_heading('Acceptance',level=1)
    para = sb_project.add_paragraph('Full Name:')
    para.add_run(user_input.firstName).bold = True
    sb_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = sb_project.add_paragraph('Date signed:')
    para.add_run(date_signature1)
    # Signature 2
    para = sb_project.add_paragraph('Full Name:')
    para.add_run(user_input.secondName).bold = True
    sb_project.add_paragraph('Signature:', style='IntenseQuote')
    # Date
    para = sb_project.add_paragraph('Date signed:')
    para.add_run(date_signature2)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=ScrumbanProject.docx'
    sb_project.save(response)
    return response

    return render(request)
